# utils/export_manager.py
# DOCX export for BRD and User Stories

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re

# ── Helpers ───────────────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def _add_paragraph(doc: Document, text: str, bold: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def _add_table_from_markdown(doc: Document, lines: list):
    """Parse markdown table lines into a Word table."""
    rows = [l for l in lines if l.startswith('|') and '---' not in l]
    if not rows:
        return
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cells)

    if not parsed:
        return

    col_count = len(parsed[0])
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = 'Table Grid'

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
            # Bold header row
            if r_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()  # spacing after table

def _parse_markdown_to_docx(doc: Document, markdown: str):
    """
    Converts markdown string to Word document content.
    Handles: # headings, | tables, bullet lists, plain paragraphs.
    """
    lines = markdown.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading
        if line.startswith('# '):
            _add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            _add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            _add_heading(doc, line[4:].strip(), 3)

        # Table — collect all consecutive table lines
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            _add_table_from_markdown(doc, table_lines)
            continue

        # Bullet
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')

        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('─' * 60)

        # Bold line **text**
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            _add_paragraph(doc, line.strip().strip('*'), bold=True)

        # Empty line — skip
        elif line.strip() == '':
            pass

        # Plain paragraph
        else:
            _add_paragraph(doc, line.strip())

        i += 1


# ── Cover Page ────────────────────────────────────────────────────────────────

def _add_cover(doc: Document, title: str, session_name: str, doc_type: str):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("🔭 ScopeCraft")
    run.font.size = Pt(28)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(doc_type).font.size = Pt(16)

    sn = doc.add_paragraph()
    sn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sn.add_run(f"Session: {session_name}").font.size = Pt(12)

    doc.add_paragraph()
    doc.add_page_break()


# ── Main Export Functions ─────────────────────────────────────────────────────

def export_brd_to_docx(brd_markdown: str, session_name: str) -> bytes:
    """
    Convert BRD markdown to DOCX bytes.
    Returns bytes for Streamlit download_button.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    _add_cover(doc, "Business Requirements Document", session_name, "Business Requirements Document")
    _parse_markdown_to_docx(doc, brd_markdown)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_stories_to_docx(stories_markdown: str, session_name: str) -> bytes:
    """
    Convert User Stories markdown to DOCX bytes.
    Returns bytes for Streamlit download_button.
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    _add_cover(doc, "User Story Backlog", session_name, "User Story Backlog — MoSCoW")
    _parse_markdown_to_docx(doc, stories_markdown)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()